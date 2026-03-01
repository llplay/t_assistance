"""
将授课计划 .docx 转换为带 {{占位符}} 的模板文件 schedule_template.docx

修复记录：
- v2: 修复封面信息行（无冒号的段落）替换逻辑
       修复多段落单元格未完全清空问题
       修复排课表模板行原内容残留问题
"""

from docx import Document
from docx.oxml.ns import qn

DOCX_SRC = "2024-2025学年第一学期授课方案(计算机视觉）23人工智能  (1).docx"
DOCX_OUT = "schedule_template.docx"

# 排课表模板行占位符（对应表格1的10列）
SCHEDULE_ROW_PLACEHOLDERS = [
    "{{月份}}", "{{周次}}", "{{星期}}", "{{节次}}",
    "{{教学内容}}", "{{教学学时}}", "{{理论}}", "{{实践}}",
    "{{重难点}}", "{{课外作业}}"
]


def clear_cell(cell, new_text: str):
    """
    清空单元格所有段落的所有 run，只在第一个段落写入 new_text。
    保留第一个段落的原有对齐等格式。
    """
    paragraphs = cell.paragraphs
    if not paragraphs:
        return

    # 第一个段落：清除所有 run，写入新内容
    p0 = paragraphs[0]
    for child in list(p0._p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "hyperlink"):
            p0._p.remove(child)
    p0.add_run(new_text)

    # 其余段落：整个从 XML 里删掉
    tc = cell._tc
    for p in paragraphs[1:]:
        tc.remove(p._p)


from docx.enum.text import WD_UNDERLINE
import os

def replace_cover_para_by_label(para, label: str, placeholder: str):
    """
    替换封面信息行中 label 之后的值部分为 placeholder。
    适用于格式 '系、部           信息工程系      '（无冒号，标签+空格+值）。
    也兼容有冒号的格式（制订日期：...）。

    策略：
      1. 按 run 遍历，累积字符确定"标签区"结束位置
      2. 标签区之后的第一个有内容的 run → 替换为 placeholder
      3. 后续有内容的 run → 清空
    """
    full_text = para.text
    if not full_text.strip():
        return

    # 找 label 在全文中的结束位置（字符偏移）
    # 优先精确匹配，兼容带冒号/不带冒号
    clean_label = label.replace("：", "").replace(":", "").replace(" ", "")
    clean_full = full_text.replace("：", "").replace(":", "").replace(" ", "")

    label_end_in_clean = clean_full.find(clean_label)
    if label_end_in_clean == -1:
        return
    label_end_in_clean += len(clean_label)

    # 映射"去除冒号和空格后的偏移"→"原文偏移"
    # 逐字符建立映射
    mapping = []  # mapping[i] = 原文 index of i-th non-space-colon char
    for idx, ch in enumerate(full_text):
        if ch not in (" ", "\t", "　", "：", ":"):
            mapping.append(idx)

    if label_end_in_clean > len(mapping):
        label_end_in_clean = len(mapping)

    # 标签在原文中的结束字符索引（含冒号/空格后缀）
    if label_end_in_clean < len(mapping):
        label_end_orig = mapping[label_end_in_clean - 1]
    else:
        label_end_orig = len(full_text)

    # 按 run 遍历，确定哪些 run 是"值区"
    char_pos = 0
    value_replaced = False
    for run in para.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        char_pos = run_end

        # 完全在标签区内 → 保留
        if run_end <= label_end_orig + 1:
            continue

        # 完全在值区内
        if run_start >= label_end_orig:
            if run.text.strip():
                if not value_replaced:
                    run.text = f"\u3000\u3000\u3000{placeholder}\u3000\u3000\u3000"
                    run.font.underline = WD_UNDERLINE.SINGLE
                    value_replaced = True
                else:
                    run.text = ""
            # 若是纯空格 run → 保留（保持视觉间距）
            else:
                pass
            continue

        # 跨越边界（标签和值在同一个 run 中）
        # 保留标签部分
        keep_part = run.text[:label_end_orig - run_start + 1]
        
        # 找到标签后面可能的空格
        i = label_end_orig - run_start + 1
        while i < len(run.text) and run.text[i] in (" ", "\t", "　"):
            keep_part += run.text[i]
            i += 1
            
        run.text = keep_part
        
        # 创建新的 run 放带下划线的值
        # 为了让 Word 强制显示下划线（特别是在行尾或空白时），我们使用中文全角空格(\u3000)进行填充补齐
        # 全角空格会被 Word 视作普通字符，100% 会渲染出连续的下划线
        padded_placeholder = f"\u3000\u3000\u3000{placeholder}\u3000\u3000\u3000" 
        
        new_run = para.add_run(padded_placeholder)
        new_run.font.underline = WD_UNDERLINE.SINGLE
        
        # 继承字体设置
        if run.font.size:
            new_run.font.size = run.font.size
        # 强制设置下划线字体颜色和字体样式，确保下划线渲染在非英文字符上也生效
        if run.font.name:
            new_run.font.name = run.font.name
            
        value_replaced = True


def make_template():
    doc = Document(DOCX_SRC)
    paras = doc.paragraphs

    # ── 1. 封面：学年学期段落（段落[3]）────────────────────
    p3 = paras[3]
    for run in p3.runs:
        run.text = ""
    if p3.runs:
        p3.runs[0].text = "（{{学年学期}}）"
    else:
        p3.add_run("（{{学年学期}}）")

    # ── 2. 封面：各信息行（无冒号格式）──────────────────────
    # (段落索引, 标签文字, 占位符)
    cover_info = [
        (4,  "系、部",        "{{系部}}"),
        (5,  "教研室",        "{{教研室}}"),
        (6,  "课程名称",      "{{课程名称}}"),
        (7,  "教学层次",      "{{教学层次}}"),
        (8,  "授课班级",      "{{授课班级}}"),
        (9,  "授课计划制订人", "{{制订人}}"),
        (10, "其他任课教师",  "{{其他任课教师}}"),
        (16, "制订日期",      "{{制订日期}}"),
    ]
    for idx, label, ph in cover_info:
        replace_cover_para_by_label(paras[idx], label, ph)

    # ── 3. 课程信息表（表格0，8行×4列）有多个段落的单元格 ──
    t0 = doc.tables[0]

    # 行0~3：值在列1、列3（非合并，单段落）
    simple_cells = [
        (0, 1, "{{课程总学时}}"),  (0, 3, "{{机动学时}}"),
        (1, 1, "{{理论教学学时}}"), (1, 3, "{{实践教学学时}}"),
        (2, 1, "{{周学时数}}"),    (2, 3, "{{教学周数}}"),
        (3, 1, "{{复习学时}}"),    (3, 3, "{{考核学时}}"),
    ]
    for ri, ci, ph in simple_cells:
        clear_cell(t0.rows[ri].cells[ci], ph)

    # 行4~6：合并单元格（教学大纲/基本教材/参考书）
    for ri, ph in [(4, "{{教学大纲}}"), (5, "{{基本教材}}"), (6, "{{参考书}}")]:
        cells = t0.rows[ri].cells
        seen_ids = set()
        for ci, cell in enumerate(cells):
            cid = id(cell._tc)
            if ci == 0:
                seen_ids.add(cid)
                continue
            if cid not in seen_ids:
                clear_cell(cell, ph)
                seen_ids.add(cid)
                break

    # 行7：课程简介（整行合并，保留标题段落，清除其余段落写占位符）
    first_cell = t0.rows[7].cells[0]
    ps = first_cell.paragraphs
    # 保留第一段（"本课程教学目的、要求及必要说明："标题），清除其余
    tc = first_cell._tc
    # 清除第2段及以后
    for p in ps[1:]:
        for child in list(p._p):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("r", "hyperlink"):
                p._p.remove(child)
        # 第一个多余段落写占位符，其余清空
    if len(ps) >= 2:
        p_target = ps[1]
        p_target.add_run("{{课程简介}}")
        # 删除第3段以后
        for p in ps[2:]:
            tc.remove(p._p)
    else:
        ps[0].add_run("\n{{课程简介}}")

    # ── 4. 排课表（表格1）：清理模板行，删除其余数据行 ───────
    t1 = doc.tables[1]
    template_row = t1.rows[2]

    # 给模板行每列写占位符（合并单元格去重）
    template_cells = template_row.cells
    seen_ids = []
    unique_cells = []
    for c in template_cells:
        cid = id(c._tc)
        if cid not in seen_ids:
            seen_ids.append(cid)
            unique_cells.append(c)

    for ci, ph in enumerate(SCHEDULE_ROW_PLACEHOLDERS):
        if ci < len(unique_cells):
            clear_cell(unique_cells[ci], ph)

    # 删除行3~末尾（倒序，避免索引位移）
    tbl_el = t1._tbl
    all_tr = tbl_el.findall(qn("w:tr"))
    for tr in reversed(all_tr[3:]):
        tbl_el.remove(tr)

    # ── 5. 保存 ──────────────────────────────────────────────
    doc.save(DOCX_OUT)
    print(f"模板已生成: {DOCX_OUT}")
    print("请用 Word 打开确认封面/信息表/排课表的占位符均正确显示。")


if __name__ == "__main__":
    make_template()
