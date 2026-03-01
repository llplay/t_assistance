from flask import Flask, request, Response, send_file
from flask_cors import CORS
import json
from docx import Document
import time
import os
import uuid
from datetime import datetime
import google.generativeai as genai
import json
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from dotenv import load_dotenv

load_dotenv()

def gen_word(gen_data):
    teaching_process_image = gen_data.get('teaching_process_image')
    gen_data = json.loads(gen_data["content"].replace("```json", "").replace("```", ''))
    # 读取教案模板
    doc = Document('教案模板.docx')
    # 遍历文档中的所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # 找到包含“##”的单元格
                if "##" in cell.text:
                    # 按照"##节点1##节点2##"的格式提取gen.json中的内容
                    keys = cell.text.split("##")  # 获取双井号之间的键
                    if "教学流程" in cell.text:
                        # 清空单元格原有内容
                        cell.text = ''
                        # 插入图片（假设图片在同一目录下）
                        image_path = teaching_process_image  # 替换为你的图片路径
                        if os.path.exists(image_path):
                            # 在文字描述后插入图片
                            picture = cell.add_paragraph()
                            run = picture.add_run()
                            run.add_picture(image_path, width=Inches(5))  # 设置图片宽度为5英寸
                            picture.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中
                    try:
                        # 根据提取的键逐层从gen.json中获取值
                        value = gen_data
                        for key in keys:
                            if key.strip() == '':
                                continue
                            # print(key)
                            value = value[key]
                        # 替换单元格内容
                        cell.text = value
                    except KeyError:
                        # 若gen.json中没有相应的键，则保持原内容不变
                        pass
                # 设置单元格字体为宋体、小四号
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
                        run.font.size = Pt(12)  # 设置小四号字体（12磅）

                # 如果是图片段落，保持居中对齐
                if "教学流程" in cell.text and os.path.exists(teaching_process_image):
                    for paragraph in cell.paragraphs:
                        if len(paragraph.runs) > 0 and hasattr(paragraph.runs[0], '_r') and paragraph.runs[0]._r.find('./{*}drawing') is not None:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 保存填充后的教案
    return doc


# 配置 Gemini
# 注意：实际使用时应通过环境变量或配置文件获取 API    # Replace with environment variable
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY not found in environment variables")
    genai.configure(api_key=api_key)
# model = genai.GenerativeModel('gemini-1.5-flash-002') # 提示：gemini-3-flash-preview 可能对应实际的 1.5 系列
model = genai.GenerativeModel('gemini-3-flash-preview') # 提示：gemini-3-flash-preview 可能对应实际的 1.5 系列


def generate_lesson_plan(data):
    subject = data["subject"]
    grade = data["grade"]
    _class = data["class"]
    students = data["students"]
    content = data["content"]
    duration = data["duration"]
    location = data["location"]
    time = data["time"]
    method = data["method"]
    teaching_process = "\n".join(data["teaching_process"])
    book_name = data["book_name"]
    example = '''
    这是以个opencv摄像头拍摄图像及视频教案json的示例
    {
    "专业名称":"计算机技术",
    "授课年级":"大二",
    "授课班级":"二班",
    "授课对象":"大专",
    "授课学时":"2",
    "授课地点":"机房",
    "授课时间":"9.8",
    "授课方式":"理实一体化",
    "教材": "",
      "授课内容": "摄像头拍摄图像及视频"
      "教学内容": "用摄像头拍摄图像及视频,本地视频的读取与播放,学习如何使用OpenCV的VideoCapture类，捕获摄像头的实时图像与视频。通过OpenCV库读取并逐帧播放本地存储的视频文件，理解视频处理流程。",
      "学情分析": {
        "知识和技能基础": "学生已经学习了Python基础语法和OpenCV的基本图像处理操作。学生对电脑摄像头有基本操作经验，了解摄像头的使用。",
        "认知和实践能力": "学生具备较强的编程理解能力，但对实时视频捕获及处理的知识还较为薄弱。学生动手能力较强，适合通过任务驱动和项目实践提升技能。",
        "学习特点": "学生动手能力强，善于通过实践操作来加深对概念的理解。学生之间具备合作学习的基础，喜欢通过讨论解决编程问题。"
      },
      "教学目标": {
        "素质目标": "培养学生的团队合作精神和自主学习能力。",
        "知识目标": "掌握摄像头图像和视频捕获的基本操作，学会读取和播放本地视频文件。",
        "能力目标": "能够使用OpenCV独立实现摄像头图像和视频捕获及本地视频播放功能。"
      },
      "教学重点": "摄像头实时图像与视频的捕获操作及本地视频的逐帧读取与播放。",
      "教学难点": "视频播放的流畅性及摄像头图像捕获的稳定性。",
      "教学方法": {
        "教法": "采用任务驱动教学法，结合课堂讲解、案例演示和实操练习。",
        "学法": "学生通过自主探究与小组合作进行编程实践，完成摄像头与视频处理任务。"
      },
      "教学过程": {
        "课前": {
          "教学环节": "课前准备",
          "教学内容": "分发学习材料",
          "教师活动": "在学习通平台上传摄像头与视频处理相关的学习材料，提供相关示例代码和操作步骤。",
          "学生活动": "学生提前预习材料，熟悉课程内容，了解摄像头与视频处理的基础知识。",
          "设计意图": "通过课前预习，帮助学生了解课程重点，提前做好准备，为课堂上的实际操作打下基础。"
        },
        "课中": {
          "明确任务": {
            "教学内容": "介绍课程任务和目标",
            "教师活动": "向学生讲解本节课的任务及目标，明确学习方向。任务包括：使用摄像头捕获图像与视频，并读取和播放本地视频。",
            "学生活动": "学生听讲并记录任务要点，明确学习目标。",
            "设计意图": "确保学生对本节课的内容有清晰的认识，能够明确自己在课堂上的任务和目标，激发学生的学习动机。"
          },
          "分析任务": {
            "教学内容": "摄像头图像与视频捕获操作",
            "教师活动": "详细讲解如何使用OpenCV的VideoCapture类来捕获摄像头的图像和视频，并现场演示代码实现的过程，解释每一步的功能和用途。",
            "学生活动": "学生听讲并做笔记，理解摄像头操作的基本流程，积极参与课堂互动，提出问题或讨论解决方法。",
            "设计意图": "通过详细讲解和现场演示，帮助学生掌握摄像头图像和视频捕获的核心技术，理解每个步骤的操作原理，为接下来的编程实践做好准备。"
          },
          "探究新知": {
            "教学内容": "视频读取与播放",
            "教师活动": "演示如何使用OpenCV读取和播放本地视频，并逐行解释代码，重点讲解视频帧的读取、显示、暂停和停止等操作。",
            "学生活动": "学生独立操作，按照教师的指导完成代码实现，尝试修改代码以实现不同的视频播放效果。",
            "设计意图": "通过演示和实际操作，帮助学生加深对视频处理的理解，培养他们独立解决问题的能力，增强课堂的互动性和参与感。"
          },
          "训练技能": {
            "教学内容": "编程实践",
            "教师活动": "布置编程实践任务，要求学生独立完成摄像头与视频处理的实际编程任务。教师在此过程中进行指导，帮助学生解决遇到的问题，并提供优化建议。",
            "学生活动": "学生独立编程，完成任务，并通过小组讨论和合作交流解决问题，共同探讨优化代码的方案。",
            "设计意图": "通过实践练习，提升学生的动手能力和解决问题的能力，同时通过合作学习，培养学生的团队合作精神和相互学习的习惯。"
          },
          "评价总结": {
            "教学内容": "课堂总结与提问",
            "教师活动": "回顾本节课的知识点，特别是摄像头与视频处理的关键操作，提问学生，检验其理解程度，并给予反馈和指导。",
            "学生活动": "学生总结课程内容，回答教师提问，主动提出自己在学习过程中遇到的难点，寻求帮助。",
            "设计意图": "通过总结和反馈帮助学生巩固知识，解决疑惑，确保每个学生都能掌握本节课的核心内容。"
          }
        },
        "课后": {
          "教学环节": "课后复习",
          "教学内容": "复习与拓展练习",
          "教师活动": "布置课后作业，要求学生完成额外的视频处理任务，如实现视频的基本编辑功能，并提供相关参考资料。",
          "学生活动": "学生完成课后练习，复习课堂所学，进一步巩固知识，并尝试独立完成拓展任务。",
          "设计意图": "通过课后作业，帮助学生在课后进一步巩固和拓展知识，强化编程实践能力，培养自主学习的习惯。"
        }
      },
      "教学评价": {
        "评价标准": "根据学生在课堂上的任务完成度、代码的正确性和优化程度以及课堂表现进行综合评估。特别关注学生独立解决问题的能力和团队合作的情况。"
      },
      "教学反思": {
        "教学效果": "大部分学生能够独立完成摄像头和视频的编程任务，且对实际应用产生浓厚兴趣，课堂互动良好，教学目标基本达成。",
        "特色创新": "结合实际应用场景，将视频处理与实时图像处理结合进行教学，并通过实际案例提升学生的学习兴趣和动手能力。",
        "不足": "部分学生在视频播放流畅性上遇到问题，缺乏优化代码的能力，对性能调优和异常处理的理解较为薄弱。",
        "改进措施": "后续课程中将加强对性能优化和异常处理的讲解，增加更多的实际案例和优化练习，帮助学生提升视频处理的代码质量和稳定性。同时，通过小组讨论和案例分析的方式，进一步培养学生的团队合作和解决问题的能力。"
      }
    }
    '''
    
    
    prompt = f'''
    我是一名高职教师，我的授课对象是{grade}{_class}{duration}{students}高职学生,上课地点是{location},上课时间是{time},
    授课方式是{method}，教材是{book_name}
    请帮我按照以下要求撰写一下教案
    这节课的主要内容有：{content}
    教案目录：
    1.教学内容
    2.学情分析
    2.1知识和技能基础
    2.2认知和实践能力
    2.3学习特点
    3.教学目标
    3.1素质目标
    3.2知识目标
    3.3能力目标
    4.教学重点
    5.教学难点
    6.教学方法
    6.1教法
    6.2学法
    8.教学过程
    8.1课前
    8.2课中
    8.3课后
    9.教学评价
    10.教学反思
    10.1教学效果
    10.2特色创新
    10.3不足
    10.4改进措施
    其中课前课后包含以下四部分内容
    教学环节
    教学内容
    教师活动
    学生活动
    设计意图
    其中课中分为以下几个部分
    {teaching_process}
    课中的每一部分也都包含教学环节、教学内容、教师活动、学生活动、设计意图
    以json的形式返回给我
    {example}
    要求内容尽可能的详细，字数要足够多， 不能少于2000字
    '''
    try:
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=4095,
                temperature=0.7,
            ),
            stream=True
        )
        for chunk in response:
            if chunk.text:
                yield chunk.text
    except Exception as e:
        yield f"\n[Generation Error: {str(e)}]"


def generate_teaching_schedule(data):
    """
    流式调用 Gemini，根据用户填写的学期信息生成授课计划 JSON。
    支持自由描述模式（前端传来的 text）或者旧的结构化表单模式。
    """
    free_text = data.get('自由描述', '').strip()

    if free_text:
        # 自由描述模式
        prompt = f"""
你是一名经验丰富的高职教师。用户用自然语言描述了一门课程的基本信息和详细要求（可能包含授课时间、教学周数、教材目录等）。
请仔细阅读以下信息，并提取相关内容生成一份完整的学期授课计划。

【用户提供的课程信息】
{free_text}

请严格按照以下 JSON 格式输出，不要加任何 markdown 标记：

{{
  "系部": "",
  "教研室": "",
  "课程名称": "",
  "教学层次": "",
  "授课班级": "",
  "制订人": "",
  "其他任课教师": "",
  "学年学期": "",
  "制订日期": "",
  "课程总学时": "",
  "机动学时": "",
  "理论教学学时": "",
  "实践教学学时": "",
  "周学时数": "",
  "教学周数": "",
  "复习学时": "",
  "考核学时": "",
  "教学大纲": "",
  "基本教材": "",
  "参考书": "",
  "课程简介": "（根据课程信息生成200字以上的教学目的、要求及必要说明）",
  "授课计划": [
    {{
      "月份": "9",
      "周次": "1",
      "星期": "4",
      "节次": "1-2",
      "教学内容": "第X章 XXXX",
      "教学学时": "2",
      "理论": "1",
      "实践": "1",
      "重难点": "重点：XXXX\\n难点：XXXX",
      "课外作业": ""
    }}
  ]
}}

生成规则：
1. 请从用户描述中提取相应的基本信息填入 JSON。无法确定的字段留空字符串。
2. 授课计划严格按教材目录或课程大纲顺序展开，覆盖要求的教学周数。
3. 每次课固定2学时，根据描述中的【上课时间】或【每周课时】排满每周的课。星期字段只填数字。
4. 如果描述中指定了上课的星期和节次，请严格按配对循环填写对应的内容。
5. 理论+实践学时之和=教学学时。
6. 重难点格式：重点：XXX\\n难点：XXX（用\\n分隔）。
7. 根据描述，在最后留出复习课和考核课时（通常也是2学时一条）。
8. 教学内容要具体，对应教材章节，不可泛泛而谈。
"""
    else:
        # 兼容旧格式（结构化表单）
        schedule_slots = data.get('上课时间', [])
        if not schedule_slots:
            weekdays_raw = data.get('上课星期', ['4', '5'])
            periods_raw  = data.get('上课节次', ['1-2'])
            wd = weekdays_raw if isinstance(weekdays_raw, list) else [weekdays_raw]
            pd_ = periods_raw  if isinstance(periods_raw,  list) else [periods_raw]
            schedule_slots = [{'星期': w, '节次': pd_[i % len(pd_)]} for i, w in enumerate(wd)]

        WEEK_MAP = {'1':'周一','2':'周二','3':'周三','4':'周四','5':'周五','6':'周六','7':'周日'}
        slots_desc = '、'.join(
            f"{WEEK_MAP.get(s.get('星期',''), '周' + s.get('星期',''))} 第{s.get('节次','')}节"
            for s in schedule_slots
        )
        # 给 AI 的结构化描述（星期数字→节次）
        slots_json_hint = ', '.join(
            f'{{"星期":"{s.get("星期","")}","节次":"{s.get("节次","")}"}}' for s in schedule_slots
        )

        prompt = f"""
你是一名经验丰富的高职教师，请根据以下课程信息，生成一份完整的学期授课计划。

课程基本信息：
- 课程名称：{data.get('课程名称', '')}
- 学年学期：{data.get('学年学期', '')}
- 授课班级：{data.get('授课班级', '')}
- 制订人：{data.get('制订人', '')}
- 课程总学时：{data.get('课程总学时', '')}（理论：{data.get('理论教学学时', '')}，实践：{data.get('实践教学学时', '')}）
- 周学时数：{data.get('周学时数', '')}，教学周数：{data.get('教学周数', '')}
- 每周上课安排：{slots_desc}（每周按此顺序循环，星期只填数字）
- 基本教材：{data.get('基本教材', '')}
- 课程简介：{data.get('课程简介', '')}

教材目录（请以此为基础安排授课内容）：
{data.get('教材目录', '')}

请严格按照以下 JSON 格式输出，不要加任何 markdown 标记：

{{
  "系部": "{data.get('系部', '')}",
  "教研室": "{data.get('教研室', '')}",
  "课程名称": "{data.get('课程名称', '')}",
  "教学层次": "{data.get('教学层次', '')}",
  "授课班级": "{data.get('授课班级', '')}",
  "制订人": "{data.get('制订人', '')}",
  "其他任课教师": "{data.get('其他任课教师', '')}",
  "学年学期": "{data.get('学年学期', '')}",
  "制订日期": "{data.get('制订日期', '')}",
  "课程总学时": "{data.get('课程总学时', '')}",
  "机动学时": "{data.get('机动学时', '')}",
  "理论教学学时": "{data.get('理论教学学时', '')}",
  "实践教学学时": "{data.get('实践教学学时', '')}",
  "周学时数": "{data.get('周学时数', '')}",
  "教学周数": "{data.get('教学周数', '')}",
  "复习学时": "{data.get('复习学时', '')}",
  "考核学时": "{data.get('考核学时', '')}",
  "教学大纲": "{data.get('教学大纲', '')}",
  "基本教材": "{data.get('基本教材', '')}",
  "参考书": "{data.get('参考书', '')}",
  "课程简介": "（根据课程信息生成200字以上的教学目的、要求及必要说明）",
  "授课计划": [
    {{
      "月份": "9",
      "周次": "1",
      "星期": "4",
      "节次": "1-2",
      "教学内容": "第X章 XXXX",
      "教学学时": "2",
      "理论": "1",
      "实践": "1",
      "重难点": "重点：XXXX\\n难点：XXXX",
      "课外作业": ""
    }}
  ]
}}

生成规则：
1. 每次课固定2学时，每周按上课时间安排生成多条记录（每个排课时间槽各一条）
2. 共需覆盖{data.get('教学周数', '18')}个教学周，授课顺序从教材目录依次展开
3. 月份从开学月份（如9月）开始，按实际月份填写
4. 每条记录的"星期"和"节次"严格按以下配对循环填写：{slots_json_hint}
   例：第1周第1条用第1个配对，第1周第2条用第2个配对，第2周第1条再用第1个配对
5. 理论+实践学时之和=教学学时
6. 重难点格式：重点：XXX\\n难点：XXX（用\\n分隔）
7. 最后留2~4节复习课，考核课时单独列
8. 教学内容要具体，对应教材章节
"""
    try:
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=32768,
                temperature=0.5,
            ),
            stream=True
        )
        for chunk in response:
            if chunk.text:
                yield chunk.text
    except Exception as e:
        yield f"\n[Generation Error: {str(e)}]"


def gen_schedule_word(schedule_json: dict) -> Document:
    """
    用 schedule_template.docx 模板填充授课计划数据，生成 Word 文档。

    模板结构：
      封面段落：{{XX}} 占位符
      表格0（课程信息，8行×4列）：{{XX}} 占位符
      表格1（排课表）：前2行为表头，第3行为模板行，复制N次填充

    参数 schedule_json: generate_teaching_schedule 生成后解析的 dict
    返回: 填充完成的 Document 对象
    """
    from copy import deepcopy
    from lxml import etree
    from docx.oxml.ns import qn as _qn

    template_path = "schedule_template.docx"
    doc = Document(template_path)

    # ── 工具：替换单元格内容（清空所有段落，写入新文本）──────
    def fill_cell(cell, value: str):
        paras = cell.paragraphs
        if not paras:
            return
        for p in paras:
            for child in list(p._p):
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag in ("r", "hyperlink"):
                    p._p.remove(child)
        lines = value.split("\n")
        paras[0].add_run(lines[0])
        for line in lines[1:]:
            paras[0].add_run("\n" + line)

    # ── 工具：段落中替换 {{占位符}} ──────────────────────────
    def fill_para(para, ph: str, val: str):
        full = para.text
        if ph not in full:
            return
        new_text = full.replace(ph, val)
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text

    # ── 1. 封面段落 ───────────────────────────────────────────
    COVER_MAP = {
        "{{学年学期}}":     schedule_json.get("学年学期", ""),
        "{{系部}}":         schedule_json.get("系部", ""),
        "{{教研室}}":       schedule_json.get("教研室", ""),
        "{{课程名称}}":     schedule_json.get("课程名称", ""),
        "{{教学层次}}":     schedule_json.get("教学层次", ""),
        "{{授课班级}}":     schedule_json.get("授课班级", ""),
        "{{制订人}}":       schedule_json.get("制订人", ""),
        "{{其他任课教师}}": schedule_json.get("其他任课教师", ""),
        "{{制订日期}}":     schedule_json.get("制订日期", ""),
    }
    for para in doc.paragraphs:
        for ph, val in COVER_MAP.items():
            if ph in para.text:
                fill_para(para, ph, val)

    # ── 2. 课程信息表（表格0）────────────────────────────────
    TABLE0_MAP = {
        "{{课程总学时}}":   schedule_json.get("课程总学时", ""),
        "{{机动学时}}":     schedule_json.get("机动学时", ""),
        "{{理论教学学时}}": schedule_json.get("理论教学学时", ""),
        "{{实践教学学时}}": schedule_json.get("实践教学学时", ""),
        "{{周学时数}}":     schedule_json.get("周学时数", ""),
        "{{教学周数}}":     schedule_json.get("教学周数", ""),
        "{{复习学时}}":     schedule_json.get("复习学时", ""),
        "{{考核学时}}":     schedule_json.get("考核学时", ""),
        "{{教学大纲}}":     schedule_json.get("教学大纲", ""),
        "{{基本教材}}":     schedule_json.get("基本教材", ""),
        "{{参考书}}":       schedule_json.get("参考书", ""),
        "{{课程简介}}":     schedule_json.get("课程简介", ""),
    }
    t0 = doc.tables[0]
    for row in t0.rows:
        seen = set()
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen:
                continue
            seen.add(cid)
            for ph, val in TABLE0_MAP.items():
                if ph in cell.text:
                    fill_cell(cell, val)
                    break

    # ── 3. 排课表（表格1）：克隆模板行逐条填充 ───────────────
    schedule_items = schedule_json.get("授课计划", [])
    if not schedule_items:
        return doc

    t1 = doc.tables[1]
    tbl_el = t1._tbl
    all_tr = tbl_el.findall(_qn("w:tr"))
    template_tr = all_tr[2]  # 模板行

    # 列字段顺序（对应模板行10列）
    FIELDS = ["月份", "周次", "星期", "节次", "教学内容",
              "教学学时", "理论", "实践", "重难点", "课外作业"]

    # 获取模板行的唯一 tc 列表（去重合并单元格）
    def unique_tc(tr):
        all_tc = tr.findall(_qn("w:tc"))
        seen, result = [], []
        for tc in all_tc:
            if id(tc) not in seen:
                seen.append(id(tc))
                result.append(tc)
        return result

    template_tcs = unique_tc(template_tr)

    # 获取模板每列第一个 run 的 rPr，用于复制字体格式
    def get_rpr(tc):
        p_list = tc.findall(_qn("w:p"))
        if not p_list:
            return None
        runs = p_list[0].findall(_qn("w:r"))
        if not runs:
            return None
        return runs[0].find(_qn("w:rPr"))

    template_rprs = [get_rpr(tc) for tc in template_tcs]

    def make_data_row(item: dict):
        new_tr = deepcopy(template_tr)
        new_tcs = unique_tc(new_tr)
        for ci, field in enumerate(FIELDS):
            if ci >= len(new_tcs):
                break
            tc = new_tcs[ci]
            p_list = tc.findall(_qn("w:p"))
            if not p_list:
                continue
            p = p_list[0]
            # 清空现有 runs
            for child in list(p):
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag in ("r", "hyperlink"):
                    p.remove(child)
            value = str(item.get(field, ""))
            # 多行内容（换行用 \n）：每行一个 run + <w:br/>
            lines = value.split("\n")
            for li, line in enumerate(lines):
                new_r = etree.SubElement(p, _qn("w:r"))
                if ci < len(template_rprs) and template_rprs[ci] is not None:
                    new_r.insert(0, deepcopy(template_rprs[ci]))
                new_t = etree.SubElement(new_r, _qn("w:t"))
                new_t.text = line
                new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                if li < len(lines) - 1:
                    br = etree.SubElement(new_r, _qn("w:br"))
        return new_tr

    # 删除模板行，按顺序插入数据行
    insert_pos = list(tbl_el).index(all_tr[1]) + 1  # 在第2行表头之后
    tbl_el.remove(template_tr)
    for item in schedule_items:
        new_tr = make_data_row(item)
        tbl_el.insert(insert_pos, new_tr)
        insert_pos += 1

    return doc

def generate_lesson_plan_for_schedule(data):
    """
    流式调用 Gemini，根据授课计划中的单节课信息及课程大纲生成一份教案 JSON。
    """
    course_info = data.get('course_info', {})
    section_info = data.get('section_info', {})
    
    prompt = f"""
    我是一名高职教师，请帮我按照以下课程信息和章节安排撰写一份详细的【教案】。
    
    【课程基本信息】
    课程名称：{course_info.get('课程名称', '')}
    授课班级：{course_info.get('授课班级', '')}
    学年学期：{course_info.get('学年学期', '')}
    课程简介：{course_info.get('课程简介', '')}
    教学层次：{course_info.get('教学层次', '')}
    
    【本节课（本次教案）信息】
    周次：第{section_info.get('周次', '')}周
    节次：第{section_info.get('节次', '')}节
    教学内容：{section_info.get('教学内容', '')}
    教学学时：{section_info.get('教学学时', '')}学时（其中理论{section_info.get('理论', '')}，实践{section_info.get('实践', '')}）
    重点难点：{section_info.get('重难点', '')}
    
    请严格按照以下 JSON格式返回，不要加任何 markdown标记：
    
    {{
      "专业名称": "{course_info.get('课程名称', '')}",
      "授课年级": "",
      "授课班级": "{course_info.get('授课班级', '')}",
      "授课对象": "{course_info.get('教学层次', '')}学生",
      "授课学时": "{section_info.get('教学学时', '')}",
      "授课地点": "多媒体教室/机房",
      "授课时间": "第{section_info.get('周次', '')}周 第{section_info.get('节次', '')}节",
      "授课方式": "理实一体化",
      "教材": "{course_info.get('基本教材', '')}",
      "授课内容": "{section_info.get('教学内容', '')}",
      "教学内容": "{section_info.get('教学内容', '')}。结合实际应用的详细描述。",
      "学情分析": {{
        "知识和技能基础": "",
        "认知和实践能力": "",
        "学习特点": ""
      }},
      "教学目标": {{
        "素质目标": "",
        "知识目标": "",
        "能力目标": ""
      }},
      "教学重点": "{section_info.get('重难点', '').split('难点：')[0].replace('重点：', '').strip() if '难点：' in section_info.get('重难点', '') else section_info.get('重难点', '')}",
      "教学难点": "{section_info.get('重难点', '').split('难点：')[1].strip() if '难点：' in section_info.get('重难点', '') else ''}",
      "教学方法": {{
        "教法": "",
        "学法": ""
      }},
      "教学过程": {{
        "课前": {{
          "教学环节": "课前准备",
          "教学内容": "",
          "教师活动": "",
          "学生活动": "",
          "设计意图": ""
        }},
        "课中": {{
          "明确任务": {{ "教学环节": "明确任务", "教学内容": "", "教师活动": "", "学生活动": "", "设计意图": "" }},
          "分析任务": {{ "教学环节": "分析任务", "教学内容": "", "教师活动": "", "学生活动": "", "设计意图": "" }},
          "探究新知": {{ "教学环节": "探究新知", "教学内容": "", "教师活动": "", "学生活动": "", "设计意图": "" }},
          "训练技能": {{ "教学环节": "训练技能", "教学内容": "", "教师活动": "", "学生活动": "", "设计意图": "" }},
          "评价总结": {{ "教学环节": "评价总结", "教学内容": "", "教师活动": "", "学生活动": "", "设计意图": "" }}
        }},
        "课后": {{
          "教学环节": "课后复习",
          "教学内容": "",
          "教师活动": "",
          "学生活动": "",
          "设计意图": ""
        }}
      }},
      "教学评价": {{
        "评价标准": ""
      }},
      "教学反思": {{
        "教学效果": "",
        "特色创新": "",
        "不足": "",
        "改进措施": ""
      }}
    }}
    
    生成要求：
    1. 结合【课程简介】和【本节课重点难点】，生成合理的学情分析、教学目标等。
    2. “课中”环节必须包含此固定五个子项，并且内部结构也必须像 JSON 中一样。
    3. 内容尽可能详细，字数要足够多，不能少于2000字。
    4. 不要包含换行符（或者如果需要换行可以用 \\n，但最后必须保证可被 JSON.parse 解析）。
    """
    
    try:
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=8192,
                temperature=0.7,
            ),
            stream=True
        )
        for chunk in response:
            if chunk.text:
                yield chunk.text
    except Exception as e:
        yield f"\\n[Generation Error: {str(e)}]"
