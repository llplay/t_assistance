"""分析 .docx 文档结构，输出到文件"""
import sys
from docx import Document

DOCX_PATH = "2024-2025学年第一学期授课方案(计算机视觉）23人工智能  (1).docx"
OUT_PATH = "docx_analysis.txt"

doc = Document(DOCX_PATH)
lines = []

lines.append("=" * 60)
lines.append("【段落】")
lines.append("=" * 60)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        lines.append(f"  [{i}] style={p.style.name!r:30s} text={p.text[:80]!r}")

lines.append("")
lines.append("=" * 60)
lines.append(f"【表格】共 {len(doc.tables)} 个")
lines.append("=" * 60)
for ti, table in enumerate(doc.tables):
    lines.append(f"\n  -- 表格 #{ti} ({len(table.rows)} 行 x {len(table.columns)} 列) --")
    for ri, row in enumerate(table.rows):
        row_texts = []
        for cell in row.cells:
            txt = cell.text.strip().replace("\n", "↵")[:40]
            row_texts.append(txt)
        lines.append(f"    行{ri:02d}: {' | '.join(row_texts)}")
        if ri >= 40:
            lines.append(f"    ... (省略剩余 {len(table.rows)-ri-1} 行)")
            break

with open(OUT_PATH, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

print(f"分析完成，结果已写入: {OUT_PATH}")
