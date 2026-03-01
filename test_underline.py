from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import zipfile
import tempfile
import xml.dom.minidom

doc = Document("schedule_template.docx")
p = doc.paragraphs[4] # 系、部

for run in p.runs:
    run.text = ""

# Add new run with underline XML directly
r_val = p.add_run("测试下划线")
r_val.font.underline = True

doc.save("test_out.docx")

# extract document.xml
with zipfile.ZipFile("test_out.docx", "r") as z:
    xml_content = z.read("word/document.xml")
    
dom = xml.dom.minidom.parseString(xml_content)
print(dom.toprettyxml(indent="  ")[:2000])
